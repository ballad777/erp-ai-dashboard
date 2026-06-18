from __future__ import annotations

from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.shared import Inches
from fastapi import UploadFile

from app.services.agent_orchestrator import (
    run_uploaded_agent_workflow,
    run_uploaded_merged_agent_workflow,
)
from app.services.artifact_access import create_artifact_url
from app.services.analysis_progress import (
    CancelCheck,
    ProgressCallback,
    emit_progress,
    ensure_not_cancelled,
)

REPO_ROOT = Path(__file__).resolve().parents[3]
GENERATED_OUTPUTS_DIR = REPO_ROOT / "generated_outputs"
REPORT_DIR = GENERATED_OUTPUTS_DIR / "reports"


async def generate_uploaded_report(
    file: UploadFile,
    target_column: str,
    analysis_mode: str = "auto",
    chart_types: str = "auto",
    model_selection_mode: str = "auto",
    selected_models: str = "auto",
    automl_mode: str = "quick",
    date_column: str | None = None,
    price_column: str | None = None,
    progress_callback: ProgressCallback | None = None,
    should_cancel: CancelCheck | None = None,
) -> dict[str, Any]:
    workflow = await run_uploaded_agent_workflow(
        file=file,
        target_column=target_column,
        analysis_mode=analysis_mode,
        chart_types=chart_types,
        model_selection_mode=model_selection_mode,
        selected_models=selected_models,
        automl_mode=automl_mode,
        date_column=date_column,
        price_column=price_column,
        progress_callback=progress_callback,
        should_cancel=should_cancel,
    )
    ensure_not_cancelled(should_cancel)
    emit_progress(progress_callback, "building_report", "正在排版並輸出 Word 報告。")
    return generate_report_from_workflow(workflow)


async def generate_uploaded_merged_report(
    files: list[UploadFile],
    target_column: str,
    analysis_mode: str = "auto",
    chart_types: str = "auto",
    model_selection_mode: str = "auto",
    selected_models: str = "auto",
    automl_mode: str = "quick",
    date_column: str | None = None,
    price_column: str | None = None,
    progress_callback: ProgressCallback | None = None,
    should_cancel: CancelCheck | None = None,
) -> dict[str, Any]:
    workflow = await run_uploaded_merged_agent_workflow(
        files=files,
        target_column=target_column,
        analysis_mode=analysis_mode,
        chart_types=chart_types,
        model_selection_mode=model_selection_mode,
        selected_models=selected_models,
        automl_mode=automl_mode,
        date_column=date_column,
        price_column=price_column,
        progress_callback=progress_callback,
        should_cancel=should_cancel,
    )
    ensure_not_cancelled(should_cancel)
    emit_progress(progress_callback, "building_report", "正在排版並輸出 Word 報告。")
    return generate_report_from_workflow(workflow)


def generate_report_from_workflow(workflow: dict[str, Any]) -> dict[str, Any]:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    report_path = REPORT_DIR / f"analysis_report_{uuid4().hex}.docx"
    decision_brief = workflow.get("decision_brief") or {}

    document = Document()
    document.add_heading("智能金融資料分析顧問報告", 0)
    document.add_paragraph("把複雜資料翻譯成一般人能理解的重點、風險、機會與下一步。")
    document.add_paragraph(f"資料檔案：{workflow['file_name']}")
    document.add_paragraph(f"目標欄位：{workflow['target_column']}")
    document.add_paragraph(f"摘要來源：{workflow['llm_provider']}")

    document.add_heading("01 執行摘要", level=1)
    for item in _safe_list(decision_brief.get("executive_summary") or workflow["executive_summary"]):
        document.add_paragraph(str(item), style="List Bullet")

    plain_summary = _safe_dict(decision_brief.get("plain_language_summary"))
    if plain_summary:
        document.add_heading("10 秒重點", level=2)
        _add_kv_paragraph(document, "發生了什麼", plain_summary.get("what_happened"))
        _add_kv_paragraph(document, "為什麼重要", plain_summary.get("why_it_matters"))
        _add_kv_paragraph(document, "主要風險", plain_summary.get("risks"))
        _add_kv_paragraph(document, "可能機會", plain_summary.get("opportunities"))
        _add_kv_paragraph(document, "下一步", plain_summary.get("next_step"))

    document.add_heading("02 決策優先級", level=1)
    priority_findings = _safe_list(decision_brief.get("priority_findings"))
    if priority_findings:
        priority_table = document.add_table(rows=1, cols=5)
        priority_table.style = "Table Grid"
        for index, header in enumerate(["類型", "重點", "說明", "依據", "建議行動"]):
            priority_table.rows[0].cells[index].text = header
        for finding in priority_findings:
            item = _safe_dict(finding)
            row = priority_table.add_row().cells
            row[0].text = str(item.get("label") or "")
            row[1].text = str(item.get("title") or "")
            row[2].text = str(item.get("summary") or "")
            row[3].text = str(item.get("evidence") or "")
            row[4].text = str(item.get("recommended_action") or "")
    else:
        document.add_paragraph("本次分析未產生決策優先級。")

    dataset = workflow["dataset_summary"]
    document.add_heading("03 資料摘要", level=1)
    document.add_paragraph(f"資料列數：{dataset['row_count']}")
    document.add_paragraph(f"欄位數：{dataset['column_count']}")
    document.add_paragraph(f"建議目標欄位：{', '.join(dataset.get('recommended_target_columns') or []) or '無'}")
    quality_report = _safe_dict(dataset.get("quality_report"))
    if quality_report:
        document.add_paragraph(f"資料品質分數：{quality_report.get('quality_score')}/100")
        for issue in _safe_list(quality_report.get("issues")):
            message = _safe_dict(issue).get("message")
            if message:
                document.add_paragraph(str(message), style="List Bullet")

    model_analysis = workflow["model_analysis"]
    document.add_heading("04 模型推薦與比較", level=1)
    document.add_paragraph(f"問題型態：{model_analysis['problem_type']}")
    document.add_paragraph(f"AutoML 模式：{model_analysis['automl_mode']}")
    model_guidance = _safe_dict(decision_brief.get("model_guidance"))
    if model_guidance:
        document.add_paragraph(str(model_guidance.get("selection_logic") or ""))
        for model in _safe_list(model_guidance.get("recommended_models"))[:6]:
            item = _safe_dict(model)
            document.add_heading(str(item.get("model_name") or "推薦模型"), level=2)
            _add_kv_paragraph(document, "模型用途", item.get("purpose"))
            _add_kv_paragraph(document, "適用資料類型", "、".join(map(str, _safe_list(item.get("suitable_data_types")))))
            _add_kv_paragraph(document, "使用難度", item.get("difficulty"))
            _add_kv_paragraph(document, "適用場景", "、".join(map(str, _safe_list(item.get("use_cases")))))
            _add_kv_paragraph(document, "為什麼推薦", item.get("why_recommended"))
            _add_kv_paragraph(document, "預期可得到", item.get("expected_output"))
            _add_kv_paragraph(document, "實際評估", item.get("evaluation_summary"))

    model_table = document.add_table(rows=1, cols=6)
    model_table.style = "Table Grid"
    headers = ["模型", "家族", "R2", "RMSE", "MAE", "訓練時間"]
    for index, header in enumerate(headers):
        model_table.rows[0].cells[index].text = header
    for metric in model_analysis["model_results"]:
        row = model_table.add_row().cells
        row[0].text = str(metric["model_name"])
        row[1].text = str(metric["model_family"])
        row[2].text = str(metric["r2"])
        row[3].text = str(metric["rmse"])
        row[4].text = str(metric["mae"])
        row[5].text = f"{metric['training_time_seconds']} 秒"

    document.add_heading("05 顧問式章節解讀", level=1)
    for section in _safe_list(decision_brief.get("report_sections")):
        _add_report_section(document, _safe_dict(section))

    financial_analysis = workflow.get("financial_analysis")
    if financial_analysis:
        document.add_heading("06 金融分析", level=1)
        for item in financial_analysis["summary"]:
            document.add_paragraph(str(item), style="List Bullet")
        document.add_paragraph(f"VaR 95%：{financial_analysis['var_95']}")
        document.add_paragraph(f"VaR 99%：{financial_analysis['var_99']}")

    document.add_heading("07 圖表解讀", level=1)
    chart_interpretations = _safe_list(decision_brief.get("chart_interpretations"))
    if not chart_interpretations:
        charts = workflow["model_analysis"]["charts"]
        if financial_analysis:
            charts = [*charts, *financial_analysis["charts"]]
        chart_interpretations = charts
    for chart in chart_interpretations[:12]:
        item = _safe_dict(chart)
        document.add_heading(str(item.get("title") or "圖表"), level=2)
        _add_kv_paragraph(document, "圖表說明", item.get("explanation"))
        _add_kv_paragraph(document, "關鍵發現", item.get("key_findings"))
        _add_kv_paragraph(document, "代表意義", item.get("meaning"))
        _add_kv_paragraph(document, "不能證明什麼", item.get("what_it_cannot_prove"))
        _add_kv_paragraph(document, "趨勢解讀", item.get("trend_interpretation"))
        _add_kv_paragraph(document, "異常說明", item.get("anomaly_note"))
        _add_kv_paragraph(document, "商業洞察", item.get("business_insight"))
        _add_kv_paragraph(document, "決策建議", item.get("recommended_action"))
        chart_path = REPO_ROOT / str(item.get("chart_path") or "")
        if chart_path.exists():
            document.add_picture(str(chart_path), width=Inches(5.8))

    document.add_heading("08 風險與限制", level=1)
    for item in _safe_list(decision_brief.get("risk_and_limitations")):
        document.add_paragraph(str(item), style="List Bullet")

    ai_conclusion = decision_brief.get("ai_conclusion")
    if ai_conclusion:
        document.add_heading("09 AI 結論摘要", level=1)
        document.add_paragraph(str(ai_conclusion))

    document.add_heading("10 分析代理執行紀錄", level=1)
    agent_table = document.add_table(rows=1, cols=3)
    agent_table.style = "Table Grid"
    hdr_cells = agent_table.rows[0].cells
    hdr_cells[0].text = "分析代理"
    hdr_cells[1].text = "狀態"
    hdr_cells[2].text = "摘要"
    for step in workflow["agent_steps"]:
        row_cells = agent_table.add_row().cells
        row_cells[0].text = str(step["agent_name"])
        row_cells[1].text = str(step["status"])
        row_cells[2].text = str(step["summary"])

    if workflow["notes"]:
        document.add_heading("11 備註", level=1)
        for note in workflow["notes"]:
            document.add_paragraph(str(note), style="List Bullet")

    document.save(report_path)

    return {
        "file_name": workflow["file_name"],
        "target_column": workflow["target_column"],
        "report_path": str(report_path.relative_to(REPO_ROOT)),
        "report_url": create_artifact_url(
            report_path,
            root=GENERATED_OUTPUTS_DIR,
        ),
        "workflow": workflow,
        "notes": [
            "已產生 Word 分析報告。",
            "一鍵 ZIP 套件目前未啟用。",
        ],
    }


def _add_report_section(document: Document, section: dict[str, Any]) -> None:
    title = str(section.get("title") or "分析章節")
    document.add_heading(title, level=2)
    _add_kv_paragraph(document, "分析目的", section.get("analysis_purpose"))
    _add_kv_paragraph(document, "分析方法", section.get("analysis_method"))
    _add_kv_paragraph(document, "分析結果", section.get("analysis_result"))
    _add_kv_paragraph(document, "結果解讀", section.get("interpretation"))
    _add_kv_paragraph(document, "商業意義", section.get("business_meaning"))
    _add_kv_paragraph(document, "建議行動", section.get("recommended_action"))
    _add_kv_paragraph(document, "風險與限制", section.get("risks_and_limits"))
    _add_kv_paragraph(document, "AI 結論摘要", section.get("ai_conclusion"))


def _add_kv_paragraph(document: Document, label: str, value: Any) -> None:
    if value is None or value == "":
        return
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}：").bold = True
    paragraph.add_run(str(value))


def _safe_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _safe_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []
